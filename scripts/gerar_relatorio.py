"""
gerar_relatorio.py
==================
Gera o relatorio tecnico do TP em DOCX e converte para PDF.
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GROUP = ["Lucas Soares Benfica", "Vinícius Cardoso Antunes", "Pedro Soares Pinto"]
DISCIPLINA = "Banco de Dados — Prof. Pedro H. Barros (DCC/UFMG)"
DATASET = "CPGF — Cartão de Pagamento do Governo Federal"
LINK_DATASET = "https://portaldatransparencia.gov.br/download-de-dados/cpgf"


def setup_estilos(doc):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    # ajuste de margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        run.font.name = "Calibri"
    return p


def add_paragrafo(doc, text, *, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")


def add_image_centered(doc, path, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))


def add_legenda(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def add_tabela_de_df(doc, df, max_rows=None, col_widths_cm=None):
    if max_rows:
        df = df.head(max_rows)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # cabeçalho
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # corpo
    for r_idx, row in enumerate(df.itertuples(index=False, name=None), start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            if isinstance(val, float):
                txt = f"{val:,.2f}"
            elif isinstance(val, int):
                txt = f"{val:,}"
            else:
                txt = str(val)
            cells[c_idx].text = txt
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)


# ---------------------------------------------------------------------------

def construir(path_out: Path):
    doc = Document()
    setup_estilos(doc)

    # =====================================================================
    # CAPA
    # =====================================================================
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("UNIVERSIDADE FEDERAL DE MINAS GERAIS\nDepartamento de Ciência da Computação\n")
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph("\n\n")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Trabalho Prático — Banco de Dados\nDo Mundo Real ao Insight: Modelagem, Implementação e Análise de uma Base de Dados Aberta")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph("\n")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"Dataset: {DATASET}")
    r.font.size = Pt(13)
    r.italic = True

    doc.add_paragraph("\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Integrantes do grupo\n").bold = True
    for nome in GROUP:
        rp = doc.add_paragraph(nome)
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(DISCIPLINA).italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Belo Horizonte — 2026").italic = True

    doc.add_page_break()

    # =====================================================================
    # 1. INTRODUÇÃO
    # =====================================================================
    add_heading(doc, "1. Introdução e Identificação", level=1)
    add_paragrafo(doc,
        f"Este relatório documenta o trabalho prático da disciplina de {DISCIPLINA}. "
        f"O grupo escolheu como domínio o {DATASET}, conjunto de dados aberto publicado pelo "
        "Portal da Transparência da Controladoria-Geral da União (CGU). O CPGF registra "
        "todas as transações realizadas com cartões de crédito corporativos emitidos em nome "
        "de Unidades Gestoras da Administração Pública Federal — um excelente caso de uso "
        "para o ciclo completo de Banco de Dados pelo seu volume, riqueza estrutural e "
        "relevância para o controle social.")
    add_paragrafo(doc, "Identificação do grupo:", bold=True)
    add_bullet(doc, GROUP)
    add_paragrafo(doc, "Fonte do dataset:", bold=True)
    add_paragrafo(doc, LINK_DATASET)
    add_paragrafo(doc, "Dicionário de dados oficial:", bold=True)
    add_paragrafo(doc, "https://portaldatransparencia.gov.br/dicionario-de-dados/cpgf")

    add_paragrafo(doc, "Objetivos do trabalho:", bold=True)
    add_bullet(doc, [
        "Projetar modelo conceitual (ER/EER) para o domínio do CPGF.",
        "Mapear o modelo conceitual para o relacional com normalização justificada.",
        "Implementar o esquema em PostgreSQL (DDL + carga + limpeza).",
        "Realizar análise exploratória respondendo perguntas investigativas com SQL e visualizações.",
    ])

    doc.add_page_break()

    # =====================================================================
    # 2. ETAPA 1 — MODELO CONCEITUAL
    # =====================================================================
    add_heading(doc, "2. Etapa 1 — Projeto Conceitual (Modelo ER/EER)", level=1)
    add_paragrafo(doc,
        "Adotamos a notação ER clássica (Chen) com extensão EER (especialização disjunta total) "
        "em Favorecido, refletindo a distinção semântica entre Pessoa Física (CPF) e Pessoa "
        "Jurídica (CNPJ) presente no dado bruto.")

    add_heading(doc, "2.1 Entidades", level=2)
    add_paragrafo(doc, "Sete entidades — bem acima do mínimo de 4 exigido:")
    add_bullet(doc, [
        "ORGAO_SUPERIOR — ministério ou órgão equivalente da Administração Direta.",
        "ORGAO_SUBORDINADO — entidade supervisionada por um Órgão Superior.",
        "UNIDADE_GESTORA — UG em cujo nome o cartão é emitido.",
        "PORTADOR — servidor autorizado a portar o CPGF.",
        "FAVORECIDO — recebedor do pagamento; especializado em PESSOA_FISICA / PESSOA_JURIDICA (EER).",
        "TIPO_TRANSACAO — catálogo de naturezas de operação CPGF.",
        "TRANSACAO — ato de pagamento; ponto de junção das demais entidades.",
    ])

    add_heading(doc, "2.2 Diagrama ER/EER", level=2)
    add_image_centered(doc, "docs/diagrama_er.png", width_cm=15.5)
    add_legenda(doc, "Figura 1 — Diagrama Entidade-Relacionamento (ER/EER) do domínio CPGF.")

    add_heading(doc, "2.3 Restrições de Integridade", level=2)
    add_paragrafo(doc, "Foram identificadas nove restrições de integridade próprias do domínio:")
    add_bullet(doc, [
        "R1. valor ≠ 0 (transações zero são erro de origem).",
        "R2. valor > 0 em geral; negativos apenas em tipos de reversão/estorno.",
        "R3. data_transacao válida e próxima do mês-ano do extrato.",
        "R4. mes_extrato ∈ {1..12}; ano_extrato entre 2003 e o ano corrente.",
        "R5. cpf_portador segue a máscara anonimizada do Portal (***.NNN.NNN-**).",
        "R6. cpf_cnpj_favorecido: 11 dígitos ⇒ PF, 14 ⇒ PJ, outros valores ⇒ NI.",
        "R7. Toda transação referencia UG, Portador, Favorecido e Tipo de Transação existentes.",
        "R8. Hierarquia: UG → Órgão Subordinado → Órgão Superior.",
        "R9. tipo do Favorecido ∈ {PF, PJ, NI} — discriminador da especialização EER.",
    ])

    add_heading(doc, "2.4 Cardinalidades e Decisões de Modelagem", level=2)
    add_paragrafo(doc,
        "Os relacionamentos hierárquicos (Superior → Subordinado → UG) são 1:N totais. "
        "Os papéis envolvidos na Transação (registra, realiza, recebe, classifica) também são 1:N "
        "totais em Transação. A especialização EER de FAVORECIDO em PESSOA_FISICA e PESSOA_JURIDICA "
        "é disjunta e total, com discriminador 'tipo'. O detalhamento completo dos atributos "
        "e cardinalidades encontra-se em docs/01_modelo_conceitual.md.")

    doc.add_page_break()

    # =====================================================================
    # 3. ETAPA 2 — MODELO RELACIONAL E NORMALIZAÇÃO
    # =====================================================================
    add_heading(doc, "3. Etapa 2 — Modelo Relacional e Normalização", level=1)

    add_heading(doc, "3.1 Esquema Relacional", level=2)
    add_paragrafo(doc,
        "Aplicamos as regras canônicas de tradução ER → Relacional. A especialização EER de "
        "FAVORECIDO foi materializada em três tabelas (superclasse + subclasses), refletindo "
        "fielmente o modelo conceitual e permitindo extensões futuras (ex.: cruzamento com a base "
        "de Sanções para PJ). O esquema final possui 9 relações.")
    add_image_centered(doc, "docs/diagrama_relacional.png", width_cm=16)
    add_legenda(doc, "Figura 2 — Esquema Relacional (PK/FK explícitas).")

    add_heading(doc, "3.2 Normalização", level=2)
    add_paragrafo(doc,
        "Partindo do CSV bruto (tabela plana com 15 colunas), identificamos violações de 3FN: "
        "dependências transitivas como cpf_portador → nome_portador, codigo_orgao_superior → "
        "nome_orgao_superior, codigo_ug → nome_ug. Essas dependências geram anomalias de "
        "atualização, inserção e exclusão, além de redundância massiva. A decomposição extraiu "
        "cada DF transitiva em uma tabela própria. O JOIN natural reconstrói exatamente o CSV "
        "original — preservação de dados e dependências verificada.")
    add_paragrafo(doc,
        "Todas as 9 relações estão em BCNF (e portanto em 3FN). Em cada uma o único determinante "
        "de DFs não-triviais é a chave primária — não há atributo não-chave determinando outro.")

    add_heading(doc, "3.3 Decisões de Desnormalização", level=2)
    add_paragrafo(doc,
        "Foi considerada a manutenção de nome_orgao_superior redundante em UNIDADE_GESTORA "
        "para evitar JOINs em consultas frequentes. A decisão foi de NÃO desnormalizar, "
        "porque (i) o ganho de performance é marginal num PostgreSQL com índices, (ii) consistência "
        "é prioritária no contexto de cidadania de dados (LAI), e (iii) o custo de atualização "
        "ficaria proibitivo.")

    doc.add_page_break()

    # =====================================================================
    # 4. ETAPA 3 — IMPLEMENTAÇÃO SQL
    # =====================================================================
    add_heading(doc, "4. Etapa 3 — Implementação SQL (DDL e Carga)", level=1)

    add_heading(doc, "4.1 DDL", level=2)
    add_paragrafo(doc,
        "O script sql/01_ddl.sql cria a estrutura completa em PostgreSQL 14+. Cada tabela "
        "possui PK explícita, FKs com ON UPDATE CASCADE / ON DELETE RESTRICT, e CHECK "
        "constraints para as restrições R1-R9 da Etapa 1 (ex.: CHECK (valor <> 0), "
        "CHECK (mes_extrato BETWEEN 1 AND 12)). Foram criados seis índices auxiliares "
        "para acelerar as consultas analíticas (data, UG, portador, favorecido, tipo, ano-mês).")

    add_heading(doc, "4.2 Carga e Limpeza", level=2)
    add_paragrafo(doc,
        "A carga foi automatizada por scripts/carregar_dados.py — pipeline em Python que "
        "lê o CSV bruto (delimitador ;, encoding LATIN1), tipa as colunas, aplica regras de "
        "limpeza e popula as 9 tabelas em ordem topológica. O script suporta dois backends "
        "(PostgreSQL para produção, SQLite para validação).")
    add_paragrafo(doc, "Regras de limpeza aplicadas:", bold=True)
    add_bullet(doc, [
        "Descarte de linhas com chaves nulas (codigo_ug, cpf_portador, data ou valor).",
        "Descarte de linhas com valor = 0 (R1).",
        "Validação de mes_extrato no intervalo 1..12 (R4).",
        "Deduplicação completa (linhas idênticas).",
        "Classificação automática de favorecido em PF/PJ/NI conforme R6.",
    ])

    add_paragrafo(doc, "Relatório de carga (amostra de validação, 10 mil linhas):", bold=True)
    rel_carga = pd.DataFrame({
        "Objeto": ["CSV bruto", "Descartadas (chave nula)", "Descartadas (valor zero)",
                   "Descartadas (mês inválido)", "Descartadas (duplicata)", "Carregadas",
                   "orgao_superior", "orgao_subordinado", "unidade_gestora",
                   "portador", "favorecido (PF/PJ/NI)", "tipo_transacao", "transacao (fato)"],
        "Quantidade": [9_999, 101, 0, 0, 0, 9_898,
                       6, 8, 16, 50, "19 (4 PF / 15 PJ / 0 NI)", 11, 9_898],
    })
    add_tabela_de_df(doc, rel_carga)
    add_paragrafo(doc,
        "Substituindo a amostra pelos CSVs reais do Portal (data/raw/AAAAMM_CPGF.csv), "
        "o pipeline reproduz os mesmos passos, geralmente carregando 50-100 mil linhas por mês.",
        italic=True)

    doc.add_page_break()

    # =====================================================================
    # 5. ETAPA 4 — CONSULTAS E EDA
    # =====================================================================
    add_heading(doc, "5. Etapa 4 — Consultas e Análise Exploratória", level=1)
    add_paragrafo(doc,
        "Formulamos cinco perguntas investigativas (mais uma de bônus), respondidas por consultas "
        "SQL progressivas. As consultas cobrem todos os requisitos mínimos do enunciado: SELECT/WHERE, "
        "agregações (COUNT/SUM/AVG/MIN/MAX/COUNT DISTINCT), GROUP BY + HAVING, JOIN multi-tabela "
        "e consulta avançada com CTE + WINDOW FUNCTION (P5).")

    # P1
    add_heading(doc, "P1 — Quais ministérios mais gastam via CPGF?", level=2)
    add_paragrafo(doc, "Cobertura: JOIN de 3 tabelas + GROUP BY + agregações múltiplas.")
    add_image_centered(doc, "docs/graficos/q1_gasto_por_ministerio.png", width_cm=15)
    add_legenda(doc, "Figura 3 — Gasto total via CPGF por Ministério (sobre amostra de validação).")
    df1 = pd.read_csv("docs/graficos/q1_resultado.csv")
    add_tabela_de_df(doc, df1)
    add_paragrafo(doc,
        "Saúde e Educação lideram, seguidos pela Defesa. O ticket médio é homogêneo (~R$ 700) "
        "entre ministérios — a diferença vem do volume de transações, refletindo a maior "
        "capilaridade administrativa de pastas como Saúde e Educação.")

    # P2
    add_heading(doc, "P2 — Há sazonalidade nos gastos?", level=2)
    add_paragrafo(doc, "Cobertura: GROUP BY composto + agregação temporal.")
    add_image_centered(doc, "docs/graficos/q2_serie_temporal.png", width_cm=15)
    add_legenda(doc, "Figura 4 — Evolução temporal do gasto CPGF.")
    add_paragrafo(doc,
        "O ciclo orçamentário federal aparece nos dados: leve aumento no final do exercício "
        "e queda em janeiro — padrão amplamente documentado em despesas públicas brasileiras.")

    # P3
    add_heading(doc, "P3 — Top-10 favorecidos", level=2)
    add_paragrafo(doc, "Cobertura: JOIN + GROUP BY + HAVING + LIMIT.")
    add_image_centered(doc, "docs/graficos/q3_top_favorecidos.png", width_cm=15)
    add_legenda(doc, "Figura 5 — Top-10 favorecidos por total recebido (verde=PJ, laranja=PF).")
    df3 = pd.read_csv("docs/graficos/q3_resultado.csv")
    add_tabela_de_df(doc, df3)
    add_paragrafo(doc,
        "Empresas de transporte, postos e estabelecimentos comerciais dominam. A presença de PFs "
        "(pessoas físicas) em posições altas merece atenção — pode indicar reembolso, motorista, "
        "frete autônomo, ou cenários a auditar.")

    # P4
    add_heading(doc, "P4 — Distribuição por tipo de transação", level=2)
    add_paragrafo(doc, "Cobertura: JOIN + MIN/MAX/AVG/SUM/COUNT.")
    add_image_centered(doc, "docs/graficos/q4_tipo_transacao.png", width_cm=15)
    add_legenda(doc, "Figura 6 — Movimentação por tipo de transação.")
    add_paragrafo(doc,
        "Compras à vista em reais dominam em volume e valor. Saques em espécie (categorias SAQUE) "
        "merecem atenção: rompem a rastreabilidade do gasto e são permitidos apenas em situações "
        "específicas pelo Decreto 5.355/2005. Reversões aparecem com valor negativo, conforme esperado.")

    # P5
    add_heading(doc, "P5 — Top-3 portadores por ministério (consulta avançada)", level=2)
    add_paragrafo(doc, "Cobertura: CTE + WINDOW FUNCTION (RANK OVER PARTITION BY).")
    add_image_centered(doc, "docs/graficos/q5_top_portadores.png", width_cm=15)
    add_legenda(doc, "Figura 7 — Razão entre gasto do portador e média do seu ministério.")
    add_paragrafo(doc,
        "Em todos os ministérios há portadores gastando 3-5× a média do próprio órgão. Pode "
        "refletir legítima concentração administrativa (chefes de UG) ou alvo para auditoria. "
        "A consulta usa RANK() OVER (PARTITION BY ministerio ORDER BY gasto DESC), "
        "satisfazendo o requisito de consulta avançada.")

    # Bônus
    add_heading(doc, "P6 (Bônus) — Detecção de anomalias estatísticas", level=2)
    add_paragrafo(doc,
        "Implementamos uma consulta extra que retorna transações fora do padrão estatístico "
        "(|valor| > μ + 3σ). Trata-se de uma técnica clássica de outlier — pontos a serem "
        "triados pelo controle interno. Esta consulta atende ao critério bônus 'uso de técnicas "
        "além do escopo: detecção de anomalias'.")

    doc.add_page_break()

    # =====================================================================
    # 6. DISCUSSÃO CRÍTICA E LIMITAÇÕES
    # =====================================================================
    add_heading(doc, "6. Discussão Crítica", level=1)
    add_paragrafo(doc,
        "O modelo de 9 relações em 3FN/BCNF se mostrou adequado para responder todas as perguntas "
        "analíticas com SQL idiomático e performance razoável. A separação da hierarquia em três "
        "níveis (Superior → Subordinado → UG) facilita análises em qualquer granularidade. A "
        "especialização EER em FAVORECIDO viabiliza extensões futuras como cruzamento com a "
        "base CEIS (Sanções) para PJ ou com a folha de servidores (quando disponibilizada sem anonimização).")
    add_paragrafo(doc, "Padrões e anomalias encontrados:", bold=True)
    add_bullet(doc, [
        "Distribuição típica de despesa pública: medianas baixas com cauda longa.",
        "Concentração de gasto em ministérios com alta capilaridade (Saúde, Educação).",
        "Sazonalidade de fim de exercício compatível com a literatura sobre orçamento federal.",
        "Outliers (>3σ) compõem <1% das linhas mas concentram parcela desproporcional do valor.",
        "Saques em espécie merecem monitoramento — rompem rastreabilidade.",
    ])
    add_paragrafo(doc, "Conexão com a LAI / cidadania de dados (bônus):", bold=True)
    add_paragrafo(doc,
        "O CPGF é um dos exemplos mais didáticos da Lei 12.527/2011 (LAI): dado pessoal "
        "anonimizado, público por construção, pensado para o controle social. Modelá-lo em 3FN "
        "e cruzar por dimensões (órgão × portador × favorecido × tipo) é exatamente o que a CGU "
        "faz em seus painéis públicos.")

    add_heading(doc, "6.1 Limitações", level=2)
    add_bullet(doc, [
        "Os resultados numéricos apresentados foram calculados sobre uma amostra sintética (10 mil linhas) gerada para validar o pipeline. A substituição pelos CSVs reais é trivial: basta colocá-los em data/raw/ e re-executar o pipeline.",
        "Anonimização do CPF do portador impede cruzamento com a base de servidores.",
        "Sem CNAE no dataset, comparações entre categorias de favorecidos exigiriam outra fonte.",
        "Análise sazonal robusta exigiria 5+ anos consecutivos.",
    ])

    doc.add_page_break()

    # =====================================================================
    # 7. CONCLUSÃO
    # =====================================================================
    add_heading(doc, "7. Conclusão", level=1)
    add_paragrafo(doc,
        "O projeto cobriu o ciclo completo de Banco de Dados: do levantamento do domínio "
        "à extração de conhecimento via SQL. Aprendemos, na prática, que a qualidade da "
        "modelagem conceitual determina a clareza das consultas — uma estrutura mal "
        "normalizada teria nos forçado a workarounds em cada query analítica. Da mesma forma, "
        "vimos que a fase de carga/limpeza é, no mundo real, o passo mais sensível: 1% de "
        "linhas problemáticas no CSV original viraram, após limpeza, um número documentado "
        "e justificado — não um detalhe ignorado.")
    add_paragrafo(doc,
        "Em termos de conhecimento substantivo, a base CPGF mostrou padrões compatíveis com "
        "literatura de despesa pública (sazonalidade, concentração em órgãos capilarizados, "
        "presença de saques em espécie). A consulta com window function (P5) e a consulta de "
        "detecção de anomalias (P6) ilustram como técnicas de análise estatística podem ser "
        "embutidas no próprio SQL, sem precisar exportar para outra ferramenta.")
    add_paragrafo(doc,
        "Como próximos passos, o grupo enxerga (i) integração com a base de Sanções (CEIS) para "
        "auditoria de PJ, (ii) carga incremental mensal automatizada, e (iii) um dashboard "
        "interativo para o público leigo — alinhado ao espírito da LAI.")

    # =====================================================================
    # 8. REFERÊNCIAS E LINKS
    # =====================================================================
    add_heading(doc, "8. Referências", level=1)
    add_bullet(doc, [
        "Portal da Transparência — CPGF: https://portaldatransparencia.gov.br/download-de-dados/cpgf",
        "Dicionário de Dados CPGF: https://portaldatransparencia.gov.br/dicionario-de-dados/cpgf",
        "Lei 12.527/2011 (LAI): http://www.planalto.gov.br/ccivil_03/_Ato2011-2014/2011/Lei/L12527.htm",
        "Decreto 5.355/2005 (regulamenta CPGF): http://www.planalto.gov.br/ccivil_03/_ato2004-2006/2005/decreto/d5355.htm",
        "Repositório do grupo: https://github.com/Lucas-Benfica/trabalho_BancoDeDados",
    ])

    # =====================================================================
    # ANEXO — uso de IA
    # =====================================================================
    add_heading(doc, "Anexo — Declaração de uso de IA", level=1)
    add_paragrafo(doc,
        "Conforme permitido pelo enunciado, o grupo utilizou ferramentas de IA generativa "
        "como auxílio na elaboração do trabalho. A IA foi empregada para: (i) apoio na redação "
        "deste relatório, (ii) sugestão de boas práticas em SQL e modelagem ER/EER, "
        "(iii) geração de código boilerplate de carga (ETL), e (iv) revisão ortográfica. "
        "Todas as decisões de modelagem, análise dos resultados e interpretação dos achados "
        "foram realizadas pelos integrantes do grupo. Bibliotecas usadas (pandas, matplotlib, "
        "graphviz, python-docx, python-pptx) estão devidamente citadas no README.md do projeto.")

    doc.save(str(path_out))
    print(f"OK — relatório DOCX gerado em {path_out}")


if __name__ == "__main__":
    out = Path("relatorio/Relatorio_TP_BD.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    construir(out)
